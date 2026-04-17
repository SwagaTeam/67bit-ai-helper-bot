"""
мок залупа не работает
"""
from mcp.server.fastmcp import FastMCP
import docx
import os
import uuid
import datetime

# Инициализируем MCP сервер
mcp = FastMCP("PromAssist_Server")

# Папка для сохранения сгенерированных документов
DOCS_DIR = "./generated_docs"
os.makedirs(DOCS_DIR, exist_ok=True)

@mcp.tool()
def generate_passport_docx(
    short_name: str, 
    full_name: str, 
    enterprise: str, 
    manager: str, 
    customer: str, 
    curator: str, 
    category: str, 
    goal: str, 
    deadline: str, 
    budget: str, 
    bonus_fund: str, 
    economic_effect: str
) -> str:
    """
    Создает Паспорт проекта в формате Word (.docx) на основе обязательных полей.
    Используй этот инструмент каждый раз, когда пользователь просит оформить паспорт или выгрузить документ проекта.
    """
    doc = docx.Document()
    doc.add_heading('ПАСПОРТ ПРОЕКТА', 0)
    
    project_id = f"PRJ-{str(uuid.uuid4())[:6].upper()}"

    fields = {
        "ID Проекта": project_id,
        "Краткое наименование": short_name,
        "Полное наименование": full_name,
        "Предприятие": enterprise,
        "Руководитель проекта": manager,
        "Заказчик": customer,
        "Куратор Проекта": curator,
        "Категория проекта": category,
        "Цель": goal,
        "Базовый срок реализации": deadline,
        "Бюджет проекта": budget,
        "Премиальный фонд команды проекта": bonus_fund,
        "Экономический эффект": economic_effect
    }

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Поле'
    hdr_cells[1].text = 'Значение'

    for key, value in fields.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value if value else "НЕ УКАЗАНО"

    filename = f"Passport_{project_id}.docx"
    filepath = os.path.join(DOCS_DIR, filename)
    doc.save(filepath)
    
    return f"Успешно! Документ Word сохранен по пути: {os.path.abspath(filepath)}"


@mcp.tool()
def generate_protocol_txt(session_date: str, participants: str, issues_discussed: str, decisions_made: str) -> str:
    """
    Создает Протокол сессии решения проблем в формате TXT.
    Вызывай, когда нужно выгрузить протокол мозгового штурма.
    """
    protocol_id = str(uuid.uuid4())[:4]
    filename = f"Protocol_{protocol_id}.txt"
    filepath = os.path.join(DOCS_DIR, filename)
    
    content = f"""ПРОТОКОЛ СЕССИИ РЕШЕНИЯ ПРОБЛЕМ
Дата: {session_date}
Участники: {participants}

1. ОБСУЖДАЕМЫЕ ПРОБЛЕМЫ:
{issues_discussed}

2. ПРИНЯТЫЕ РЕШЕНИЯ (по методикам 6С / SMED / TPM):
{decisions_made}
"""
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content.strip())

    return f"Протокол успешно сгенерирован и сохранен: {os.path.abspath(filepath)}"


@mcp.tool()
def register_idea_jira(title: str, description: str, author: str) -> str:
    """
    Интеграция с Фабрикой Идей (Jira SM).
    Регистрирует новую идею в системе и возвращает номер тикета.
    Вызывай, когда пользователь просит "зарегистрировать идею", "отправить в джиру", "сохранить на фабрику".
    """
    # Имитация API-запроса в Jira
    ticket_id = f"IDEA-{datetime.datetime.now().strftime('%M%S')}"
    
    # В реальной жизни здесь был бы requests.post("https://jira.company.com/rest/api/...")
    
    response = (
        f"Идея '{title}' от '{author}' успешно зарегистрирована в Jira SM! "
        f"Номер тикета: {ticket_id}. Статус: Ожидает оценки отдела улучшений."
    )
    return response

if __name__ == "__main__":
    # Запуск сервера на стандартном потоке ввода/вывода (stdio), как того требует MCP
    mcp.run()
