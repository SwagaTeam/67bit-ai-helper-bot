import os
import uuid
import datetime 
import docx
from weasyprint import HTML 
from jinja2 import Environment, FileSystemLoader
from mcp.server.fastmcp import FastMCP
from atlassian import Jira

JIRA_URL = os.getenv("JIRA_URL")
JIRA_EMAIL = os.getenv("JIRA_EMAIL")
JIRA_TOKEN = os.getenv("JIRA_TOKEN") 

jira = Jira(url=JIRA_URL, username=JIRA_EMAIL, password=JIRA_TOKEN, cloud=True)

mcp = FastMCP("PromAssist_Service", stateless_http=True, host='0.0.0.0', port=8000)
env = Environment(loader=FileSystemLoader('.'))
DOCS_DIR = "/app/generated_docs"
os.makedirs(DOCS_DIR, exist_ok=True) 

@mcp.tool()
def register_idea_jira(short_name: str, full_name: str, enterprise: str, project_manager: str, project_key: str = "PROJ") -> str:
    """
    Регистрация идеи в Jira: создание основной задачи и 4 подзадач (стандарт 6.2).
    """
    try:
        summary = f"[Фабрика идей] {short_name}"
        description = f"Полное наименование: {full_name}\nПредприятие: {enterprise}\nРП: {project_manager}"
        
        # 1. Создание родительской задачи
        parent_issue = jira.issue_create(fields={
            'project': {'key': project_key},
            'summary': summary,
            'description': description,
            'issuetype': {'name': 'Task'}
        })
        parent_key = parent_issue['key']

        # 2. Список подзадач согласно регламенту
        subtasks = [
            {"summary": "Оценка в отделе улучшений", "desc": "Назначено: Отдел улучшений"},
            {"summary": "Детализация проекта и расчет эффекта", "desc": "Подготовка ТЭО"},
            {"summary": "Сбор согласований", "desc": "Владелец процесса (Директор)"},
            {"summary": "Утверждение", "desc": "Фин. директор, Тех. директор"}
        ]

        # 3. Создание подзадач
        created_subtasks = []
        for st in subtasks:
            child = jira.issue_create(fields={
                'project': {'key': project_key},
                'summary': st['summary'],
                'description': st['desc'],
                'issuetype': {'name': 'Sub-task'},
                'parent': {'key': parent_key}
            })
            created_subtasks.append(child['key'])

        return f"Задачи созданы в Jira. Родитель: {parent_key}. Подзадачи: {', '.join(created_subtasks)}"
    except Exception as e:
        return f"Ошибка Jira: {str(e)}"

@mcp.tool()
def generate_passport_pdf(
    short_name: str, full_name: str, enterprise: str, project_manager: str,
    customer: str, curator: str, category: str, goal: str, 
    deadline: str, budget: str, bonus_fund: str, economic_effect: str
) -> str:
    """Генерирует PDF паспорт проекта (стандарт 6.1)."""
    try:
        template = env.get_template('template.html')
        
        # Подготовка данных для рендеринга
        render_data = {
            "short_name": short_name, "full_name": full_name, "enterprise": enterprise,
            "project_manager": project_manager, "customer": customer, "curator": curator,
            "category": category, "goal": goal, "deadline": deadline,
            "budget": budget, "bonus_fund": bonus_fund, "economic_effect": economic_effect,
            "now": datetime.datetime.now().strftime("%d.%m.%Y %H:%M")
        }
        
        html_out = template.render(**render_data)
        filename = f"Passport_{uuid.uuid4().hex[:6]}.pdf"
        file_path = os.path.join(DOCS_DIR, filename)
        
        HTML(string=html_out).write_pdf(file_path)
        return f"PDF создан: {filename}"
    except Exception as e:
        return f"Ошибка PDF: {str(e)}"

@mcp.tool()
def generate_passport_docx(
    short_name: str, full_name: str, enterprise: str, project_manager: str,
    customer: str, curator: str, category: str, goal: str, 
    deadline: str, budget: str, bonus_fund: str, economic_effect: str
) -> str:
    """Генерирует Word паспорт (.docx) со всеми полями (стандарт 6.1)."""
    try:
        doc = docx.Document()
        doc.add_heading('ПАСПОРТ ПРОЕКТА', 0)
        
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        data_map = [
            ("Краткое наименование", short_name),
            ("Полное наименование", full_name),
            ("Предприятие", enterprise),
            ("Руководитель проекта", project_manager),
            ("Заказчик", customer),
            ("Куратор", curator),
            ("Категория", category),
            ("Цель", goal),
            ("Срок реализации", deadline),
            ("Бюджет", budget),
            ("Премиальный фонд", bonus_fund),
            ("Экономический эффект", economic_effect)
        ]
        
        for key, value in data_map:
            row = table.add_row().cells
            row[0].text = key
            row[1].text = str(value) if value else "—"

        filename = f"Passport_{uuid.uuid4().hex[:6]}.docx"
        file_path = os.path.join(DOCS_DIR, filename)
        doc.save(file_path)
        return f"Word документ создан: {filename}"
    except Exception as e:
        return f"Ошибка Word: {str(e)}"

if __name__ == "__main__":
    mcp.run(transport='streamable-http')